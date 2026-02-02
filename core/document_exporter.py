"""
Document exporter module for creating .docx files with articles and questions.
"""

from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from io import BytesIO
from typing import List, Dict


def create_article_document(article: str, questions: List[Dict]) -> BytesIO:
    """
    Create a Word document with article and questions.

    Args:
        article: The article text
        questions: List of question dictionaries with 'question' and 'type' keys

    Returns:
        BytesIO object containing the .docx file
    """
    # Create document
    doc = Document()

    # Set default font
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Arial'
    font.size = Pt(12)

    # Add title
    title = doc.add_heading('English Reading Practice', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Add article section
    doc.add_heading('Article', level=1)
    article_para = doc.add_paragraph(article)
    article_para.alignment = WD_ALIGN_PARAGRAPH.LEFT
    # Make article text larger and easier to read
    for run in article_para.runs:
        run.font.size = Pt(14)
        run.font.name = 'Arial'

    # Add some spacing
    doc.add_paragraph()

    # Add questions section
    doc.add_heading('Comprehension Questions', level=1)

    for i, q in enumerate(questions, 1):
        # Question number and type
        question_heading = doc.add_heading(f'Question {i} ({q["type"]})', level=2)

        # Question text
        doc.add_paragraph(q['question'])

        # Add options for multiple choice questions
        if q.get('type') == 'multiple_choice' and 'options' in q:
            options_para = doc.add_paragraph()
            options_para.add_run('Options:').bold = True
            for option in q['options']:
                doc.add_paragraph(option, style='List Bullet')

        # Answer space
        doc.add_paragraph('Answer:')
        doc.add_paragraph('_' * 80)

        # Add spacing between questions
        if i < len(questions):
            doc.add_paragraph()

    # Save to BytesIO
    file_stream = BytesIO()
    doc.save(file_stream)
    file_stream.seek(0)

    return file_stream


def create_article_with_answers_document(article: str, questions: List[Dict],
                                         answers: List[str], evaluation: Dict) -> BytesIO:
    """
    Create a Word document with article, questions, answers, and evaluation.

    Args:
        article: The article text
        questions: List of question dictionaries
        answers: List of user answers
        evaluation: Evaluation dictionary with score and feedback

    Returns:
        BytesIO object containing the .docx file
    """
    # Create document
    doc = Document()

    # Set default font
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Arial'
    font.size = Pt(12)

    # Add title
    title = doc.add_heading('English Reading Practice - Results', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Add score
    score_para = doc.add_paragraph()
    score_para.add_run(f'Score: {evaluation["score"]}/100').bold = True
    score_para.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Add article section
    doc.add_heading('Article', level=1)
    article_para = doc.add_paragraph(article)
    article_para.alignment = WD_ALIGN_PARAGRAPH.LEFT
    # Make article text larger and easier to read
    for run in article_para.runs:
        run.font.size = Pt(14)
        run.font.name = 'Arial'

    # Add spacing
    doc.add_paragraph()

    # Add questions and answers section
    doc.add_heading('Questions and Answers', level=1)

    for i, (q, ans) in enumerate(zip(questions, answers), 1):
        # Question number and type
        doc.add_heading(f'Question {i} ({q["type"]})', level=2)

        # Question text
        doc.add_paragraph(q['question'])

        # Add options for multiple choice questions
        if q.get('type') == 'multiple_choice' and 'options' in q:
            options_para = doc.add_paragraph()
            options_para.add_run('Options:').bold = True
            for option in q['options']:
                doc.add_paragraph(option, style='List Bullet')

        # User answer
        ans_para = doc.add_paragraph()
        ans_para.add_run('Your Answer: ').bold = True
        ans_para.add_run(ans or '(No answer provided)')

        # Show correct answer
        correct_ans_para = doc.add_paragraph()
        correct_ans_para.add_run('Correct Answer: ').bold = True
        correct_ans_para.add_run(str(q.get('correct_answer', 'N/A')))

        # Feedback
        if i <= len(evaluation.get('item_analysis', [])):
            analysis = evaluation['item_analysis'][i-1]
            correct = analysis.get('correct', False)
            feedback = analysis.get('feedback', 'N/A')

            feedback_para = doc.add_paragraph()
            status_text = '✓ Correct' if correct else '✗ Incorrect'
            status_run = feedback_para.add_run(status_text + ': ')
            status_run.bold = True
            feedback_para.add_run(feedback)

        # Add spacing between questions
        if i < len(questions):
            doc.add_paragraph()

    # Add overall feedback
    doc.add_heading('Overall Feedback', level=1)
    doc.add_paragraph(evaluation.get('overall_feedback', 'N/A'))

    # Add suggestions
    doc.add_heading('Suggestions', level=1)
    doc.add_paragraph(evaluation.get('suggestions', 'N/A'))

    # Save to BytesIO
    file_stream = BytesIO()
    doc.save(file_stream)
    file_stream.seek(0)

    return file_stream
